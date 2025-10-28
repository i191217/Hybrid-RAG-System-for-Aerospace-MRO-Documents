#!/usr/bin/env python3
"""
Complete document processing pipeline.
Filters documents, creates embeddings, and stores in Qdrant.
"""

import sys
import os
import time
import logging
import uuid
import hashlib
import io
from pathlib import Path
from typing import List, Dict, Any
import json
import numpy as np

# Add current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Import our modules
from config import config
from document_filter import DocumentFilter, FileInfo
from embedding_service import EmbeddingService
from vector_db import create_vector_database, VectorPoint

# PDF and OCR processing imports
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False

# Setup logging
logger = logging.getLogger("doc_processor.pipeline")

class DocumentProcessingPipeline:
    """Complete document processing pipeline."""
    
    def __init__(self):
        """Initialize the document processing pipeline."""
        self.logger = logging.getLogger("doc_processor.pipeline.DocumentProcessingPipeline")
        
        try:
            # Initialize components
            self.vector_db = None
            self.embedding_service = EmbeddingService()
            self.document_filter = DocumentFilter()
            
            # Initialize EasyOCR reader if available
            self.easyocr_reader = None
            if EASYOCR_AVAILABLE:
                try:
                    self.easyocr_reader = easyocr.Reader(['en'])
                    self.logger.info("EasyOCR reader initialized successfully")
                except Exception as e:
                    self.logger.warning(f"Failed to initialize EasyOCR reader: {e}")
            else:
                self.logger.warning("EasyOCR not available. PDF text extraction will use fallback methods.")
            
                # Initialize statistics
            self.stats = {
                    "files_found": 0,
                "files_processed": 0,
                "embeddings_created": 0,
                "embeddings_stored": 0,
                    "errors": 0
            }
            
            # Initialize debug file
            self._initialize_debug_file()
        
            self.logger.info("Initialized DocumentProcessingPipeline")
            
        except Exception as e:
            self.logger.error(f"Failed to initialize pipeline: {e}")
            raise
    
    def _initialize_debug_file(self):
        """Initialize the chunks debug file."""
        try:
            with open("chunks.txt", "w", encoding="utf-8") as debug_file:
                debug_file.write("DOCUMENT PROCESSING PIPELINE - CHUNKS DEBUG\n")
                debug_file.write("=" * 100 + "\n")
                debug_file.write(f"Started at: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                debug_file.write("=" * 100 + "\n\n")
            self.logger.info("Initialized chunks debug file: chunks.txt")
        except Exception as e:
            self.logger.error(f"Failed to initialize debug file: {e}")
    
    def initialize_vector_db(self):
        """Initialize the vector database connection."""
        try:
            self.logger.info(f"Initializing vector database: {config.VECTOR_DB_TYPE}")
            
            # Create vector database instance
            self.vector_db = create_vector_database(config.COLLECTION_NAME)
            
            # Check if collection exists and delete it if configured
            collection_exists = self.vector_db.collection_exists()
            if collection_exists:
                self.logger.info(f"Collection {config.COLLECTION_NAME} already exists")
                if config.CLEAR_COLLECTION:
                    self.logger.info(f"CLEAR_COLLECTION=true, deleting existing collection: {config.COLLECTION_NAME}")
                    self.vector_db.delete_collection()
                    collection_exists = False
                else:
                    self.logger.info(f"CLEAR_COLLECTION=false, keeping existing collection")
            
            # Create collection if it doesn't exist
            if not collection_exists:
                self.logger.info(f"Creating collection: {config.COLLECTION_NAME}")
                self.vector_db.create_collection(
                    dimension=config.EMBEDDING_DIMENSION
                )
                self.logger.info(f"Created collection with dimension: {config.EMBEDDING_DIMENSION}")
            
            # Get collection info
            info = self.vector_db.get_collection_info()
            self.logger.info(f"Collection info: {info}")
            
            return True
            
        except Exception as e:
            self.logger.error(f"Error initializing vector database: {e}")
            return False
    
    def filter_documents(self, input_directory: str) -> List[FileInfo]:
        """
        Filter documents in the input directory.
        
        Args:
            input_directory: Path to directory containing documents
            
        Returns:
            List of FileInfo objects for files that passed filtering
        """
        self.logger.info(f"Starting document filtering for: {input_directory}")
        
        input_path = Path(input_directory)
        if not input_path.exists():
            raise ValueError(f"Input directory does not exist: {input_directory}")
        
        # Process directory with filtering
        file_infos = self.document_filter.process_directory(input_path)
        
        # Separate kept and rejected files
        kept_files = [f for f in file_infos if not f.reject_reason]
        rejected_files = [f for f in file_infos if f.reject_reason]
        
        # Update statistics
        self.stats["files_found"] = len(file_infos)
        self.stats["files_processed"] = len(kept_files)
        
        # Log results
        self.logger.info(f"Filtering complete:")
        self.logger.info(f"  Total files found: {len(file_infos)}")
        self.logger.info(f"  Files kept: {len(kept_files)}")
        self.logger.info(f"  Files rejected: {len(rejected_files)}")
        
        if rejected_files:
            self.logger.info("Rejected files:")
            for file_info in rejected_files[:5]:  # Show first 5
                self.logger.info(f"  - {file_info.filename}: {file_info.reject_reason}")
            if len(rejected_files) > 5:
                self.logger.info(f"  ... and {len(rejected_files) - 5} more")
        
        return kept_files
    
    def extract_text_from_file(self, file_path: Path) -> str:
        """
        Extract text from a file.
        
        Supports:
        - PDFs using PyPDF2
        - Excel files using openpyxl
        - Word documents using python-docx
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        self.logger.info(f"EXTRACTING TEXT FROM: {file_path.name}")
        self.logger.info(f"Full path: {file_path}")
        self.logger.info(f"File size: {file_path.stat().st_size / (1024*1024):.2f} MB")
        
        filename = file_path.name
        extension = file_path.suffix.lower()
        self.logger.info(f"File type: {extension}")
        
        try:
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_from_word(file_path)
            elif extension == '.txt':
                return self._extract_from_text(file_path)
            else:
                self.logger.warning(f"Unsupported file type: {extension}")
                return f"Unsupported file type: {extension}"
                
        except Exception as e:
            self.logger.error(f"Failed to extract text from {filename}: {e}")
            return f"Error extracting text from {filename}: {str(e)}"
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files using EasyOCR and PyMuPDF."""
        self.logger.info(f"Extracting PDF content using EasyOCR...")
        
        try:
            # Check if required libraries are available
            if not PYMUPDF_AVAILABLE:
                self.logger.error("PyMuPDF not available. Install: pip install PyMuPDF")
                return self._extract_from_pdf_fallback(file_path)
            
            if not EASYOCR_AVAILABLE or not self.easyocr_reader:
                self.logger.warning("EasyOCR not available. Using fallback method...")
                return self._extract_from_pdf_fallback(file_path)
            
            if not PIL_AVAILABLE:
                self.logger.error("PIL not available. Install: pip install Pillow")
                return self._extract_from_pdf_fallback(file_path)
            
            # Convert PDF to images using PyMuPDF
            self.logger.info(f"Using PyMuPDF + EasyOCR...")
            images = self._pdf_to_images_pymupdf(file_path)
            
            if not images:
                self.logger.warning("Failed to convert PDF to images. Using fallback method...")
                return self._extract_from_pdf_fallback(file_path)
            
            # Extract text from each page using EasyOCR
            page_texts = []
            failed_pages = []
            
            # Create progress bar for page processing
            if TQDM_AVAILABLE:
                page_iterator = tqdm(enumerate(images), total=len(images), 
                                   desc=f"Processing {file_path.name} with EasyOCR", 
                                   unit="page", 
                                   leave=False)
            else:
                page_iterator = enumerate(images)
            
            for page_num, image in page_iterator:
                try:
                    self.logger.debug(f"      Processing page {page_num + 1} with EasyOCR...")
                    
                    # Convert PIL image to numpy array for EasyOCR
                    img_array = np.array(image)
                    
                    # Perform OCR
                    results = self.easyocr_reader.readtext(img_array, detail=0)
                    
                    # Join all text results
                    text = ' '.join(results)
                    
                    if text.strip():
                        page_texts.append(f"=== PAGE {page_num + 1} ===\n{text}\n")
                        self.logger.debug(f"      Page {page_num + 1}: {len(text)} characters extracted")
                    else:
                        self.logger.warning(f"      Page {page_num + 1}: No text extracted")
                        failed_pages.append(page_num + 1)
                        page_texts.append(f"=== PAGE {page_num + 1} ===\n[No text extracted]\n")
                
                except Exception as e:
                    self.logger.warning(f"      Page {page_num + 1}: EasyOCR failed - {e}")
                    failed_pages.append(page_num + 1)
                    page_texts.append(f"=== PAGE {page_num + 1} ===\n[OCR extraction failed: {e}]\n")
            
            # Combine all page texts
            combined_text = '\n'.join(page_texts)
            
            self.logger.info(f"   PDF extraction complete using EasyOCR:")
            self.logger.info(f"      - Total pages: {len(images)}")
            self.logger.info(f"      - Successfully processed: {len(images) - len(failed_pages)}")
            self.logger.info(f"      - Failed pages: {len(failed_pages)}")
            self.logger.info(f"      - Total characters: {len(combined_text)}")
            
            return combined_text
            
        except Exception as e:
            self.logger.error(f"  EasyOCR PDF extraction failed: {e}")
            self.logger.info("   Attempting fallback method...")
            return self._extract_from_pdf_fallback(file_path)
    
    def _pdf_to_images_pymupdf(self, pdf_path: Path, dpi: int = 300) -> List:
        """Convert PDF pages to images using PyMuPDF."""
        try:
            self.logger.debug(f"      Converting PDF to images using PyMuPDF...")
            
            # Open PDF with PyMuPDF
            doc = fitz.open(str(pdf_path))
            images = []
            
            # Convert each page to image
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Create transformation matrix for DPI
                mat = fitz.Matrix(dpi / 72, dpi / 72)
                
                # Render page as image
                pix = page.get_pixmap(matrix=mat)
                
                # Convert to PIL Image
                img_data = pix.tobytes("ppm")
                image = Image.open(io.BytesIO(img_data))
                images.append(image)
            
            doc.close()
            self.logger.debug(f"      Converted {len(images)} pages to images")
            return images
            
        except Exception as e:
            self.logger.error(f"      PyMuPDF PDF to image conversion failed: {e}")
            return []
    
    def _extract_from_pdf_fallback(self, file_path: Path) -> str:
        """Fallback PDF extraction using traditional methods."""
        self.logger.info(f"   Using fallback PDF extraction methods...")
        
        try:
            # Try PyPDF2 first
            import PyPDF2
            self.logger.info(f"   Using PyPDF2...")
            
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                self.logger.info(f"   PDF has {len(pdf_reader.pages)} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                            self.logger.debug(f"      Page {page_num + 1}: {len(page_text)} chars")
                        else:
                            self.logger.warning(f"      Page {page_num + 1}: No text extracted")
                    except Exception as e:
                        self.logger.warning(f"      Page {page_num + 1}: Error - {e}")
                        continue
            
            combined_text = '\n'.join(text_content)
            self.logger.info(f"   PDF extraction complete: {len(combined_text)} characters from {len(text_content)} pages")
            return combined_text
            
        except ImportError:
            self.logger.warning("   PyPDF2 not available, trying pdfplumber...")
            try:
                import pdfplumber
                self.logger.info(f"   Using pdfplumber...")
                
                text_content = []
                with pdfplumber.open(file_path) as pdf:
                    self.logger.info(f"   PDF has {len(pdf.pages)} pages")
                    
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_content.append(page_text)
                                self.logger.debug(f"      Page {page_num + 1}: {len(page_text)} chars")
                            else:
                                self.logger.warning(f"      Page {page_num + 1}: No text extracted")
                        except Exception as e:
                            self.logger.warning(f"      Page {page_num + 1}: Error - {e}")
                            continue
                
                combined_text = '\n'.join(text_content)
                self.logger.info(f"   PDF extraction complete: {len(combined_text)} characters from {len(text_content)} pages")
                return combined_text
                
            except ImportError:
                self.logger.error("   No PDF libraries available. Install: pip install PyPDF2 or pip install pdfplumber")
                return f"PDF extraction failed: No PDF libraries available for {file_path.name}"
        
        except Exception as e:
            self.logger.error(f"   PDF extraction failed: {e}")
            return f"PDF extraction failed: {str(e)}"
    
    def _extract_from_excel(self, file_path: Path) -> str:
        """Extract text from Excel files."""
        self.logger.info(f"   Extracting Excel content...")
        
        try:
            import openpyxl
            self.logger.info(f"   Using openpyxl...")
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            self.logger.info(f"   Excel has {len(workbook.sheetnames)} sheets: {workbook.sheetnames}")
            
            all_text = []
            
            for sheet_name in workbook.sheetnames:
                self.logger.info(f"   Processing sheet: {sheet_name}")
                sheet = workbook[sheet_name]
                
                sheet_text = []
                sheet_text.append(f"=== SHEET: {sheet_name} ===")
                
                # Get the used range
                if sheet.max_row > 0 and sheet.max_column > 0:
                    self.logger.info(f"      Range: {sheet.max_row} rows x {sheet.max_column} columns")
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = []
                        for cell_value in row:
                            if cell_value is not None:
                                row_text.append(str(cell_value))
                        
                        if row_text:  # Only add non-empty rows
                            sheet_text.append('\t'.join(row_text))
                    
                    self.logger.info(f"      Extracted {len(sheet_text)-1} non-empty rows")
                else:
                    self.logger.warning(f"      Sheet {sheet_name} appears to be empty")
                
                if len(sheet_text) > 1:  # More than just the header
                    all_text.extend(sheet_text)
                    all_text.append("")  # Add spacing between sheets
            
            combined_text = '\n'.join(all_text)
            self.logger.info(f"   Excel extraction complete: {len(combined_text)} characters")
            return combined_text
            
        except ImportError:
            self.logger.warning("   openpyxl not available, trying pandas...")
            try:
                import pandas as pd
                self.logger.info(f"   Using pandas...")
                
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                self.logger.info(f"   Excel has {len(excel_file.sheet_names)} sheets: {excel_file.sheet_names}")
                
                all_text = []
                
                for sheet_name in excel_file.sheet_names:
                    self.logger.info(f"   Processing sheet: {sheet_name}")
                    try:
                        df = pd.read_excel(file_path, sheet_name=sheet_name)
                        
                        if not df.empty:
                            sheet_text = [f"=== SHEET: {sheet_name} ==="]
                            
                            # Add column headers
                            headers = [str(col) for col in df.columns if str(col) != 'nan']
                            if headers:
                                sheet_text.append('\t'.join(headers))
                            
                            # Add data rows
                            for _, row in df.iterrows():
                                row_values = [str(val) for val in row.values if str(val) != 'nan' and val is not None]
                                if row_values:
                                    sheet_text.append('\t'.join(row_values))
                            
                            all_text.extend(sheet_text)
                            all_text.append("")  # Add spacing
                            self.logger.info(f"      Extracted {len(df)} rows x {len(df.columns)} columns")
                        else:
                            self.logger.warning(f"      Sheet {sheet_name} is empty")
                    
                    except Exception as e:
                        self.logger.warning(f"      Failed to read sheet {sheet_name}: {e}")
                        continue
                
                combined_text = '\n'.join(all_text)
                self.logger.info(f"   Excel extraction complete: {len(combined_text)} characters")
                return combined_text
                
            except ImportError:
                self.logger.error("   No Excel libraries available. Install: pip install openpyxl or pip install pandas")
                return f"Excel extraction failed: No Excel libraries available for {file_path.name}"
        
        except Exception as e:
            self.logger.error(f"   Excel extraction failed: {e}")
            return f"Excel extraction failed: {str(e)}"
    
    def _extract_from_word(self, file_path: Path) -> str:
        """Extract text from Word documents."""
        self.logger.info(f"   Extracting Word document content...")
        
        try:
            import docx
            self.logger.info(f"   Using python-docx...")
            
            doc = docx.Document(file_path)
            
            text_content = []
            
            # Extract paragraphs
            self.logger.info(f"   Document has {len(doc.paragraphs)} paragraphs")
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    
            # Extract tables
            if doc.tables:
                self.logger.info(f"   Document has {len(doc.tables)} tables")
                for table_num, table in enumerate(doc.tables):
                    text_content.append(f"\n=== TABLE {table_num + 1} ===")
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append('\t'.join(row_text))
            
            combined_text = '\n'.join(text_content)
            self.logger.info(f"   Word extraction complete: {len(combined_text)} characters")
            return combined_text
            
        except ImportError:
            self.logger.error("  python-docx not available. Install: pip install python-docx")
            return f"Word extraction failed: python-docx library not available for {file_path.name}"
        
        except Exception as e:
            self.logger.error(f"   Word extraction failed: {e}")
            return f"Word extraction failed: {str(e)}"
    
    def _extract_from_text(self, file_path: Path) -> str:
        """Extract text from plain text files."""
        self.logger.info(f"  Reading text file...")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                self.logger.info(f"  Text file read: {len(content)} characters")
                return content
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin1') as file:
                    content = file.read()
                    self.logger.info(f"   Text file read (latin1): {len(content)} characters")
                    return content
            except Exception as e:
                self.logger.error(f"   Failed to read text file: {e}")
                return f"Text extraction failed: {str(e)}"
    
    def create_document_chunks(self, file_info: FileInfo, text_content: str) -> List[Dict[str, Any]]:
        """
        Create chunks from document text.
        
        Args:
            file_info: FileInfo object
            text_content: Extracted text content
            
        Returns:
            List of document chunks with metadata
        """
        # DEBUG: Save original text content to debug file
        with open("chunks.txt", "a", encoding="utf-8") as debug_file:
            debug_file.write("=" * 100 + "\n")
            debug_file.write(f"PROCESSING FILE: {file_info.filename}\n")
            debug_file.write(f"FILE SIZE: {file_info.size / (1024*1024):.2f} MB\n")
            debug_file.write(f"EXTRACTED TEXT LENGTH: {len(text_content)} characters\n")
            debug_file.write("-" * 50 + "\n")
            debug_file.write("EXTRACTED TEXT CONTENT:\n")
            debug_file.write("-" * 50 + "\n")
            if text_content.strip():
                debug_file.write(text_content[:2000])  # First 2000 chars
                if len(text_content) > 2000:
                    debug_file.write(f"\n... [TRUNCATED - Total length: {len(text_content)} chars]")
            else:
                debug_file.write("[NO TEXT CONTENT EXTRACTED]")
            debug_file.write("\n" + "-" * 50 + "\n")
        
        # Chunking configuration
        chunk_size = 500  # characters
        self.logger.info(f"    CHUNKING: target size = {chunk_size} characters")
        
        # Simple chunking strategy - split by sentences and group
        sentences = text_content.split('.')
        chunks = []
        
        current_chunk = ""
        chunk_index = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                # Generate UUID for this chunk
                chunk_uuid = self.generate_point_id(current_chunk.strip(), file_info.filename, chunk_index)
                
                # Save current chunk
                chunk_data = {
                    "chunk_id": chunk_uuid,
                    "content": current_chunk.strip(),
                    "filename": file_info.filename,
                    "file_path": str(file_info.path),
                    "chunk_index": chunk_index,
                    "file_size": file_info.size,
                    "file_extension": file_info.extension,
                    "metadata": {
                        "source": file_info.filename,
                        "chunk_index": chunk_index,
                        "file_size_mb": round(file_info.size / (1024 * 1024), 2),
                        "processing_timestamp": time.time(),
                        "chunk_uuid": chunk_uuid,
                        "content": current_chunk.strip()  # Include content in metadata for vector DB
                    }
                }
                chunks.append(chunk_data)
                
                # DEBUG: Save chunk to debug file
                with open("chunks.txt", "a", encoding="utf-8") as debug_file:
                    debug_file.write(f"CHUNK {chunk_index + 1}:\n")
                    debug_file.write(f"  Length: {len(current_chunk.strip())} characters\n")
                    debug_file.write(f"  Content: {current_chunk.strip()}\n")
                    debug_file.write("-" * 30 + "\n")
                
                current_chunk = sentence
                chunk_index += 1
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add the last chunk if it exists
        if current_chunk.strip():
            # Generate UUID for the final chunk
            chunk_uuid = self.generate_point_id(current_chunk.strip(), file_info.filename, chunk_index)
            
            chunk_data = {
                "chunk_id": chunk_uuid,
                "content": current_chunk.strip(),
                "filename": file_info.filename,
                "file_path": str(file_info.path),
                "chunk_index": chunk_index,
                "file_size": file_info.size,
                "file_extension": file_info.extension,
                "metadata": {
                    "source": file_info.filename,
                    "chunk_index": chunk_index,
                    "file_size_mb": round(file_info.size / (1024 * 1024), 2),
                    "processing_timestamp": time.time(),
                    "chunk_uuid": chunk_uuid,
                    "content": current_chunk.strip()  # Include content in metadata for vector DB
                }
            }
            chunks.append(chunk_data)
            
            # DEBUG: Save final chunk to debug file
            with open("chunks.txt", "a", encoding="utf-8") as debug_file:
                debug_file.write(f"CHUNK {chunk_index + 1} (FINAL):\n")
                debug_file.write(f"  Length: {len(current_chunk.strip())} characters\n")
                debug_file.write(f"  Content: {current_chunk.strip()}\n")
                debug_file.write("-" * 30 + "\n")
        
        # Log chunking results
        if chunks:
            chunk_sizes = [len(chunk["content"]) for chunk in chunks]
            avg_size = sum(chunk_sizes) / len(chunk_sizes)
            self.logger.info(f"     CHUNKING COMPLETE: {len(chunks)} chunks created")
            self.logger.info(f"       Sizes: avg={avg_size:.0f}, min={min(chunk_sizes)}, max={max(chunk_sizes)} chars")
        else:
            self.logger.warning(f"    No chunks created from {file_info.filename}")
        
        # DEBUG: Save summary to debug file
        with open("chunks.txt", "a", encoding="utf-8") as debug_file:
            debug_file.write(f"SUMMARY FOR {file_info.filename}:\n")
            debug_file.write(f"  Original text: {len(text_content)} characters\n")
            debug_file.write(f"  Created chunks: {len(chunks)}\n")
            debug_file.write(f"  Total chunk content: {sum(len(chunk['content']) for chunk in chunks)} characters\n")
            if chunks:
                chunk_sizes = [len(chunk["content"]) for chunk in chunks]
                debug_file.write(f"  Average chunk size: {sum(chunk_sizes) / len(chunk_sizes):.0f} characters\n")
                debug_file.write(f"  Chunk size range: {min(chunk_sizes)} - {max(chunk_sizes)} characters\n")
            debug_file.write("=" * 100 + "\n\n")
        
        return chunks
    
    def process_documents(self, kept_files: List[FileInfo]) -> bool:
        """
        Process filtered documents to create embeddings.
        
        Args:
            kept_files: List of FileInfo objects that passed filtering
            
        Returns:
            True if processing succeeded, False otherwise
        """
        self.logger.info(f"Starting document processing for {len(kept_files)} files")
        self.logger.info(f"BATCH SIZE: {config.BATCH_SIZE} chunks per batch")
        
        total_chunks = 0
        processed_chunks = 0
        
        try:
            for i, file_info in enumerate(kept_files, 1):
                self.logger.info(f"Processing file {i}/{len(kept_files)}: {file_info.filename}")
                
                try:
                    # Extract text from file
                    text_content = self.extract_text_from_file(file_info.path)
                    
                    # Create chunks
                    chunks = self.create_document_chunks(file_info, text_content)
                    total_chunks += len(chunks)
                    
                    self.logger.info(f"  Created {len(chunks)} chunks from {file_info.filename}")
                    
                    # Show chunk size statistics
                    if chunks:
                        chunk_sizes = [len(chunk["content"]) for chunk in chunks]
                        avg_size = sum(chunk_sizes) / len(chunk_sizes)
                        self.logger.info(f"     Chunk sizes: avg={avg_size:.0f}, min={min(chunk_sizes)}, max={max(chunk_sizes)} characters")
                    
                    # Process chunks in batches
                    batch_size = config.BATCH_SIZE
                    total_batches = (len(chunks) + batch_size - 1) // batch_size
                    self.logger.info(f"  Processing {len(chunks)} chunks in {total_batches} batches of {batch_size}")
                    
                    for batch_num in range(total_batches):
                        batch_start = batch_num * batch_size
                        batch_end = min(batch_start + batch_size, len(chunks))
                        batch_chunks = chunks[batch_start:batch_end]
                        
                        self.logger.info(f"    Processing batch {batch_num + 1}/{total_batches}: chunks {batch_start + 1}-{batch_end}")
                        
                        # Create embeddings for batch
                        texts = [chunk["content"] for chunk in batch_chunks]
                        embeddings = self.embedding_service.embed_batch(texts)
                        
                        if embeddings:
                            # Create vector points
                            vector_points = []
                            for j, (chunk, embedding) in enumerate(zip(batch_chunks, embeddings)):
                                point = VectorPoint(
                                    id=self.generate_point_id(chunk["content"], chunk["filename"], chunk["chunk_index"]),
                                    vector=embedding,
                                    metadata=chunk["metadata"]
                                )
                                vector_points.append(point)
                            
                            # Store in vector database
                            success = self.vector_db.upsert_points(vector_points)
                            
                            if success:
                                processed_chunks += len(batch_chunks)
                                self.stats["embeddings_created"] += len(embeddings)
                                self.stats["embeddings_stored"] += len(vector_points)
                                self.logger.info(f"   Stored batch {batch_num + 1}: {len(batch_chunks)} chunks successfully")
                            else:
                                self.logger.error(f"    Failed to store batch {batch_num + 1} for {file_info.filename}")
                                self.stats["errors"] += 1
                        else:
                            self.logger.error(f"      Failed to create embeddings for batch {batch_num + 1} in {file_info.filename}")
                            self.stats["errors"] += 1
                    
                    self.stats["files_processed"] += 1
                    
                except Exception as e:
                    self.logger.error(f"Error processing file {file_info.filename}: {e}")
                    self.stats["errors"] += 1
                    continue
            
            # Log final statistics
            self.logger.info(f"Document processing complete:")
            self.logger.info(f"  Files processed: {self.stats['files_processed']}/{len(kept_files)}")
            self.logger.info(f"  Total chunks created: {total_chunks}")
            self.logger.info(f"  Chunks successfully processed: {processed_chunks}")
            self.logger.info(f"  Embeddings created: {self.stats['embeddings_created']}")
            self.logger.info(f"  Embeddings stored: {self.stats['embeddings_stored']}")
            self.logger.info(f"  Errors: {self.stats['errors']}")
            
            if total_chunks > 0:
                success_rate = processed_chunks / total_chunks
                self.logger.info(f"  Processing success rate: {success_rate:.1%}")
            
            # Return success only if we processed at least some files and had minimal errors
            success_rate = self.stats['files_processed'] / len(kept_files) if kept_files else 0
            if success_rate < 0.5:  # Less than 50% success rate
                self.logger.error(f"Processing failed: only {success_rate:.1%} of files were successfully processed")
                return False
            
            return True
            
        except Exception as e:
            self.logger.error(f"Error during document processing: {e}")
            return False
    
    def run_pipeline(self, input_directory: str) -> bool:
        """
        Run the complete document processing pipeline.
        
        Args:
            input_directory: Directory containing documents to process
            
        Returns:
            True if pipeline succeeded, False otherwise
        """
        self.stats["start_time"] = time.time()
        
        try:
            self.logger.info("=" * 60)
            self.logger.info("STARTING Document Processing Pipeline")
            self.logger.info("=" * 60)
            
            # Step 1: Initialize vector database
            self.logger.info("Step 1: Initializing vector database...")
            if not self.initialize_vector_db():
                self.logger.error("Failed to initialize vector database")
                return False
            
            # Step 2: Filter documents
            self.logger.info("Step 2: Filtering documents...")
            kept_files = self.filter_documents(input_directory)
            
            if not kept_files:
                self.logger.warning("No files passed filtering. Pipeline complete.")
                return True
            
            # Step 3: Process documents
            self.logger.info("Step 3: Processing documents and creating embeddings...")
            success = self.process_documents(kept_files)
            
            # Calculate final statistics
            self.stats["end_time"] = time.time()
            total_time = self.stats["end_time"] - self.stats["start_time"]
            
            self.logger.info("=" * 60)
            self.logger.info("PIPELINE COMPLETE - Final Statistics")
            self.logger.info("=" * 60)
            self.logger.info(f"Total execution time: {total_time:.2f} seconds")
            self.logger.info(f"Files found: {self.stats['files_found']}")
            self.logger.info(f"Files processed: {self.stats['files_processed']}")
            self.logger.info(f"Embeddings created: {self.stats['embeddings_created']}")
            self.logger.info(f"Embeddings stored: {self.stats['embeddings_stored']}")
            self.logger.info(f"Errors encountered: {self.stats['errors']}")
            
            if self.stats['embeddings_created'] > 0:
                self.logger.info(f"Average embedding time: {total_time/self.stats['embeddings_created']:.3f}s per embedding")
            
            return success
            
        except Exception as e:
            self.logger.error(f"Pipeline failed with error: {e}")
            import traceback
            traceback.print_exc()
            return False
        
        finally:
            # Cleanup
            if self.vector_db:
                self.vector_db.close()
            self.embedding_service.log_stats()

    def generate_point_id(self, chunk_content: str, filename: str, chunk_index: int) -> str:
        """
        Generate a consistent UUID for a chunk based on its content and metadata.
        
        Args:
            chunk_content: The text content of the chunk
            filename: The source filename
            chunk_index: The index of the chunk within the file
            
        Returns:
            A UUID string that will be consistent for the same content
        """
        # Create a deterministic hash from content + metadata
        hash_input = f"{filename}_{chunk_index}_{chunk_content}".encode('utf-8')
        content_hash = hashlib.md5(hash_input).hexdigest()
        
        # Convert hash to UUID format (deterministic UUID)
        uuid_str = f"{content_hash[:8]}-{content_hash[8:12]}-{content_hash[12:16]}-{content_hash[16:20]}-{content_hash[20:32]}"
        
        return uuid_str

def main():
    """Main function to run the pipeline."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Document Processing Pipeline")
    parser.add_argument(
        "input_directory", 
        help="Directory containing documents to process"
    )
    parser.add_argument(
        "--test", 
        action="store_true", 
        help="Run in test mode (process only a few files)"
    )
    
    args = parser.parse_args()
    
    # Update config if in test mode
    if args.test:
        config.TEST_MODE = True
        config.BATCH_SIZE = min(config.BATCH_SIZE, 10)
        logger.info("Running in TEST MODE")
    
    # Validate configuration
    if not config.validate():
        logger.error("Configuration validation failed. Please check your .env file.")
        return False
    
    # Check if input directory exists
    if not Path(args.input_directory).exists():
        logger.error(f"Input directory does not exist: {args.input_directory}")
        return False
    
    # Run pipeline
    pipeline = DocumentProcessingPipeline()
    success = pipeline.run_pipeline(args.input_directory)
    
    if success:
        logger.info("SUCCESS: Pipeline completed successfully!")
    else:
        logger.error("ERROR: Pipeline failed. Check the logs for details.")
    
    return success

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 